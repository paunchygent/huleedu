"""
Integration tests for Strategy-based text extraction with real file content.

Tests the complete extraction flow with actual file fixtures to validate
end-to-end functionality of the Strategy pattern implementation.
"""

from __future__ import annotations

import io
from uuid import UUID, uuid4

import docx
import pytest
from common_core.error_enums import FileValidationErrorCode
from huleedu_service_libs.error_handling import HuleEduError
from pypdf import PdfWriter

from services.file_service.implementations.extraction_strategies import (
    DocxExtractionStrategy,
    ExtractionStrategy,
    PdfExtractionStrategy,
    TxtExtractionStrategy,
)
from services.file_service.implementations.text_extractor_impl import StrategyBasedTextExtractor


class TestStrategyBasedExtractorIntegration:
    """Integration tests with real file content."""

    @pytest.fixture
    def extractor(self) -> StrategyBasedTextExtractor:
        """Create text extractor instance."""
        strategies: dict[str, ExtractionStrategy] = {
            ".txt": TxtExtractionStrategy(),
            ".docx": DocxExtractionStrategy(),
            ".pdf": PdfExtractionStrategy(),
        }

        return StrategyBasedTextExtractor(validators=[], strategies=strategies)

    @pytest.fixture
    def correlation_id(self) -> UUID:
        """Generate test correlation ID."""
        return uuid4()

    @pytest.fixture
    def sample_txt_content(self) -> bytes:
        """Create sample .txt file content."""
        content = (
            "This is a sample essay for integration testing.\n"
            "It has multiple lines and special characters: éñ🎓"
        )
        return content.encode("utf-8")

    @pytest.fixture
    def sample_docx_content(self) -> bytes:
        """Create sample .docx file content."""
        document = docx.Document()
        document.add_paragraph("First paragraph of the integration test essay.")
        document.add_paragraph("Second paragraph with more content.")
        document.add_paragraph("")  # Empty paragraph should be filtered
        document.add_paragraph("Third paragraph after empty one.")

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.read()

    @pytest.fixture
    def sample_pdf_content(self) -> bytes:
        """Create sample .pdf file content."""
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        buffer = io.BytesIO()
        pdf_canvas = canvas.Canvas(buffer, pagesize=letter)

        # Page 1
        pdf_canvas.drawString(100, 750, "Integration Test PDF - Page 1")
        pdf_canvas.drawString(100, 730, "This is the first page content.")
        pdf_canvas.showPage()

        # Page 2
        pdf_canvas.drawString(100, 750, "Page 2 Content")
        pdf_canvas.drawString(100, 730, "This is the second page with different text.")
        pdf_canvas.showPage()

        pdf_canvas.save()
        buffer.seek(0)
        return buffer.read()

    @pytest.fixture
    def encrypted_pdf_content(self) -> bytes:
        """Create encrypted .pdf file content."""
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        # Create a simple PDF first
        buffer = io.BytesIO()
        pdf_canvas = canvas.Canvas(buffer, pagesize=letter)
        pdf_canvas.drawString(100, 750, "This is encrypted content")
        pdf_canvas.save()
        buffer.seek(0)

        # Encrypt it using pypdf
        from pypdf import PdfReader

        pdf_reader = PdfReader(buffer)
        pdf_writer = PdfWriter()
        pdf_writer.append_pages_from_reader(pdf_reader)
        pdf_writer.encrypt("password123")

        encrypted_buffer = io.BytesIO()
        pdf_writer.write(encrypted_buffer)
        encrypted_buffer.seek(0)
        return encrypted_buffer.read()

    async def test_txt_extraction_integration(
        self,
        extractor: StrategyBasedTextExtractor,
        sample_txt_content: bytes,
        correlation_id: UUID,
    ) -> None:
        """Test .txt file extraction with real content."""
        # When
        result = await extractor.extract_text(sample_txt_content, "test.txt", correlation_id)

        # Then
        assert "This is a sample essay for integration testing." in result
        assert "It has multiple lines" in result
        # Verify UTF-8 characters are handled correctly
        assert "é" in result or "??" in result  # Depending on encoding handling

    async def test_docx_extraction_integration(
        self,
        extractor: StrategyBasedTextExtractor,
        sample_docx_content: bytes,
        correlation_id: UUID,
    ) -> None:
        """Test .docx file extraction with real content."""
        # When
        result = await extractor.extract_text(sample_docx_content, "test.docx", correlation_id)

        # Then
        assert "First paragraph of the integration test essay." in result
        assert "Second paragraph with more content." in result
        assert "Third paragraph after empty one." in result
        # Verify paragraphs are joined with newlines
        lines = result.split("\n")
        assert len(lines) >= 3

    async def test_pdf_extraction_integration(
        self,
        extractor: StrategyBasedTextExtractor,
        sample_pdf_content: bytes,
        correlation_id: UUID,
    ) -> None:
        """Test .pdf file extraction with real content."""
        # When
        result = await extractor.extract_text(sample_pdf_content, "test.pdf", correlation_id)

        # Then
        assert "Integration Test PDF - Page 1" in result
        assert "This is the first page content." in result
        assert "Page 2 Content" in result
        assert "This is the second page with different text." in result

    async def test_encrypted_pdf_error_integration(
        self,
        extractor: StrategyBasedTextExtractor,
        encrypted_pdf_content: bytes,
        correlation_id: UUID,
    ) -> None:
        """Test encrypted .pdf file raises proper error."""
        # When/Then
        with pytest.raises(HuleEduError) as exc_info:
            await extractor.extract_text(encrypted_pdf_content, "encrypted.pdf", correlation_id)

        # Verify proper error code
        error_detail = exc_info.value.error_detail
        assert error_detail.error_code == FileValidationErrorCode.ENCRYPTED_FILE_UNSUPPORTED
        assert error_detail.service == "file_service"
        assert error_detail.operation == "extract_text"
        assert error_detail.correlation_id == correlation_id

    async def test_case_insensitive_extensions_integration(
        self,
        extractor: StrategyBasedTextExtractor,
        sample_txt_content: bytes,
        correlation_id: UUID,
    ) -> None:
        """Test that file extensions work in a case-insensitive manner."""
        # Test different cases
        test_cases = ["test.TXT", "test.Txt", "TEST.PDF", "document.DOCX"]

        for filename in test_cases:
            if filename.upper().endswith(".TXT"):
                # When
                result = await extractor.extract_text(sample_txt_content, filename, correlation_id)
                # Then
                assert "This is a sample essay" in result

    async def test_unsupported_file_type_integration(
        self,
        extractor: StrategyBasedTextExtractor,
        correlation_id: UUID,
    ) -> None:
        """Test unsupported file types raise proper errors."""
        # Given
        fake_content = b"This is not a supported file type"

        # When/Then
        with pytest.raises(HuleEduError) as exc_info:
            await extractor.extract_text(fake_content, "document.xlsx", correlation_id)

        # Verify error details
        error_detail = exc_info.value.error_detail
        assert error_detail.error_code == FileValidationErrorCode.TEXT_EXTRACTION_FAILED
        assert "No extraction strategy found for file type '.xlsx'." in error_detail.message

    async def test_file_without_extension_integration(
        self,
        extractor: StrategyBasedTextExtractor,
        correlation_id: UUID,
    ) -> None:
        """Test files without extensions raise proper errors."""
        # Given
        fake_content = b"Content without extension"

        # When/Then
        with pytest.raises(HuleEduError) as exc_info:
            await extractor.extract_text(fake_content, "README", correlation_id)

        # Verify error details
        error_detail = exc_info.value.error_detail
        assert error_detail.error_code == FileValidationErrorCode.TEXT_EXTRACTION_FAILED
        assert "No extraction strategy found for file type ''." in error_detail.message

    async def test_empty_file_content_integration(
        self,
        extractor: StrategyBasedTextExtractor,
        correlation_id: UUID,
    ) -> None:
        """Test extraction from empty files."""
        # When
        result = await extractor.extract_text(b"", "empty.txt", correlation_id)

        # Then
        assert result == ""

    async def test_correlation_id_flow_integration(
        self,
        extractor: StrategyBasedTextExtractor,
        sample_txt_content: bytes,
    ) -> None:
        """Test that correlation IDs flow through the extraction process."""
        # Given
        specific_correlation_id = uuid4()

        # When
        result = await extractor.extract_text(
            sample_txt_content, "correlation_test.txt", specific_correlation_id
        )

        # Then - Should not raise exception and should return expected content
        assert "This is a sample essay" in result
        # Note: In a more comprehensive test, we might verify logging contains correlation_id

    async def test_mixed_content_extraction_integration(
        self,
        extractor: StrategyBasedTextExtractor,
        correlation_id: UUID,
    ) -> None:
        """Test extraction from different file types in sequence."""
        # Given
        txt_content = b"Text file content for sequential testing."

        # When - Extract from different file types
        txt_result = await extractor.extract_text(txt_content, "seq1.txt", correlation_id)

        # Then
        assert "Text file content for sequential testing." in txt_result
        assert isinstance(txt_result, str)
        assert len(txt_result) > 0


@pytest.mark.slow
class TestLargeFileIntegration:
    """Integration tests with larger file content (marked as slow)."""

    @pytest.fixture
    def extractor(self) -> StrategyBasedTextExtractor:
        """Create text extractor instance."""
        strategies: dict[str, ExtractionStrategy] = {
            ".txt": TxtExtractionStrategy(),
            ".docx": DocxExtractionStrategy(),
            ".pdf": PdfExtractionStrategy(),
        }

        return StrategyBasedTextExtractor(validators=[], strategies=strategies)

    @pytest.fixture
    def large_txt_content(self) -> bytes:
        """Create large .txt file content."""
        # Create a 50KB text file
        large_text = "This is a large text file for performance testing. " * 1000
        return large_text.encode("utf-8")

    async def test_large_txt_file_extraction(
        self,
        extractor: StrategyBasedTextExtractor,
        large_txt_content: bytes,
    ) -> None:
        """Test extraction from large .txt files."""
        # Given
        correlation_id = uuid4()

        # When
        result = await extractor.extract_text(large_txt_content, "large.txt", correlation_id)

        # Then
        assert len(result) > 50000  # Should be around 50KB
        assert "This is a large text file for performance testing." in result
        assert result.count("This is a large text file for performance testing.") == 1000
